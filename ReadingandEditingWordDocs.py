import docx, os

os.chdir("c:\\users\\burha\\MyPythonScripts\\SampleFiles")
print(os.getcwd())

d = docx.Document('demo.docx')

# returns document object
d.paragraphs

# document objects contain paragraphs member variable
# paragraph objects contain text member variable

print(d.paragraphs[0].text)
print(d.paragraphs[1].text)


p = d.paragraphs[1]
print(p.runs)

# each para contains a list of runs; runs also have text member variable
# a new run is characterised by a change in style

print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

# runs also have bold,italic,underline member variables

p.runs[3].underline = True
p.runs[3].text = 'italic and underlined'

d.save('c:\\users\\burha\\MyPythonScripts\\SampleFiles\\doc2.docx')

# once done editing use save method to save the file

# press ctrl+alt+shift+s in word to bring up word styles
# paragraph and run objects also have style member variables

p.style = 'Title'

d.save('c:\\users\\burha\\MyPythonScripts\\SampleFiles\\demo3.docx')




## creating new doc

d = docx.Document()

d.add_paragraph('Hello this is a paragraph')

d.add_paragraph('This is another paragraph')


d.save('c:\\users\\burha\\MyPythonScripts\\SampleFiles\\demo4.doc')


p = d.paragraphs[0]

p.add_run('This is a new run')

p.runs[1].bold = True

d.save('c:\\users\\burha\\MyPythonScripts\\SampleFiles\\demo5.doc')



# func to get all text from a doc


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('C:\\users\\burha\\MyPythonScripts\\SampleFiles\\demo.docx'))






